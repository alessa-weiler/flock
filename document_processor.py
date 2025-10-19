"""
Document Processor for extracting text from various file formats
Supports: PDF (with OCR), DOCX, TXT, MD, CSV
"""

import os
import io
import csv as csv_module
from typing import Dict, List, Optional
from datetime import datetime

# PDF processing
from PyPDF2 import PdfReader

# DOCX processing
from docx import Document as DocxDocument

# OCR for scanned PDFs
try:
    import pytesseract
    from PIL import Image
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR not available. Install pytesseract and pdf2image for scanned PDF support")


class DocumentExtractor:
    """Extracts text and metadata from various document formats"""

    def __init__(self):
        self.ocr_available = OCR_AVAILABLE

    def extract(self, file_path: str, file_type: str) -> Dict:
        """
        Main extraction method - routes to appropriate handler

        Args:
            file_path: Path to the file
            file_type: Type of file (pdf, docx, txt, md, csv)

        Returns:
            Dict with text, metadata, and structure info
        """
        extractors = {
            'pdf': self.extract_pdf,
            'docx': self.extract_docx,
            'txt': self.extract_text,
            'md': self.extract_text,
            'csv': self.extract_csv
        }

        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")

        return extractor(file_path)

    def extract_pdf(self, file_path: str) -> Dict:
        """
        Extract text from PDF, with OCR fallback for scanned documents

        Args:
            file_path: Path to PDF file

        Returns:
            Dict with text, metadata, and page info
        """
        try:
            reader = PdfReader(file_path)

            # Extract metadata
            metadata = {
                'page_count': len(reader.pages),
                'author': None,
                'created': None,
                'modified': None
            }

            # Try to get PDF metadata
            if reader.metadata:
                metadata['author'] = reader.metadata.get('/Author', None)
                if '/CreationDate' in reader.metadata:
                    try:
                        metadata['created'] = reader.metadata['/CreationDate']
                    except:
                        pass
                if '/ModDate' in reader.metadata:
                    try:
                        metadata['modified'] = reader.metadata['/ModDate']
                    except:
                        pass

            # Extract text from all pages
            pages = []
            full_text = []

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                pages.append({
                    'page_number': i + 1,
                    'text': page_text,
                    'char_count': len(page_text)
                })
                full_text.append(page_text)

            combined_text = '\n\n'.join(full_text)

            # Check if OCR is needed (very little text extracted)
            avg_chars_per_page = len(combined_text) / len(reader.pages) if reader.pages else 0

            if avg_chars_per_page < 100 and self.ocr_available:
                print(f"PDF appears to be scanned (avg {avg_chars_per_page:.0f} chars/page), attempting OCR...")
                ocr_result = self.extract_pdf_with_ocr(file_path)
                if ocr_result:
                    return ocr_result

            return {
                'text': combined_text,
                'metadata': metadata,
                'pages': pages,
                'extraction_method': 'pypdf2'
            }

        except Exception as e:
            raise Exception(f"Failed to extract PDF: {str(e)}")

    def extract_pdf_with_ocr(self, file_path: str) -> Optional[Dict]:
        """
        Extract text from scanned PDF using OCR

        Args:
            file_path: Path to PDF file

        Returns:
            Dict with extracted text or None if OCR fails
        """
        if not self.ocr_available:
            return None

        try:
            # Convert PDF to images
            images = pdf2image.convert_from_path(file_path)

            pages = []
            full_text = []

            for i, image in enumerate(images):
                # Perform OCR on each page
                page_text = pytesseract.image_to_string(image)
                pages.append({
                    'page_number': i + 1,
                    'text': page_text,
                    'char_count': len(page_text)
                })
                full_text.append(page_text)

            combined_text = '\n\n'.join(full_text)

            return {
                'text': combined_text,
                'metadata': {
                    'page_count': len(images),
                    'author': None,
                    'created': None,
                    'modified': None
                },
                'pages': pages,
                'extraction_method': 'ocr'
            }

        except Exception as e:
            print(f"OCR extraction failed: {e}")
            return None

    def extract_docx(self, file_path: str) -> Dict:
        """
        Extract text from DOCX preserving structure

        Args:
            file_path: Path to DOCX file

        Returns:
            Dict with text, structure, and metadata
        """
        try:
            doc = DocxDocument(file_path)

            # Extract metadata from core properties
            metadata = {
                'author': doc.core_properties.author,
                'created': doc.core_properties.created.isoformat() if doc.core_properties.created else None,
                'modified': doc.core_properties.modified.isoformat() if doc.core_properties.modified else None,
                'title': doc.core_properties.title
            }

            # Extract text with structure
            paragraphs = []
            tables = []
            full_text = []

            for paragraph in doc.paragraphs:
                para_text = paragraph.text
                if para_text.strip():
                    para_data = {
                        'text': para_text,
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    }
                    paragraphs.append(para_data)
                    full_text.append(para_text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    'table_number': table_idx + 1,
                    'data': table_data
                })

                # Add table to text as formatted string
                table_text = '\n'.join([' | '.join(row) for row in table_data])
                full_text.append(f"\n[Table {table_idx + 1}]\n{table_text}\n")

            combined_text = '\n'.join(full_text)

            return {
                'text': combined_text,
                'metadata': metadata,
                'structure': {
                    'paragraphs': paragraphs,
                    'tables': tables,
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables)
                },
                'extraction_method': 'python-docx'
            }

        except Exception as e:
            raise Exception(f"Failed to extract DOCX: {str(e)}")

    def extract_text(self, file_path: str) -> Dict:
        """
        Extract text from plain text or markdown files

        Args:
            file_path: Path to text file

        Returns:
            Dict with text and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                raise Exception("Could not decode file with any supported encoding")

            # Count lines and characters
            lines = text.split('\n')

            metadata = {
                'line_count': len(lines),
                'char_count': len(text),
                'word_count': len(text.split()),
                'encoding': encoding
            }

            return {
                'text': text,
                'metadata': metadata,
                'extraction_method': 'plain_text'
            }

        except Exception as e:
            raise Exception(f"Failed to extract text file: {str(e)}")

    def extract_csv(self, file_path: str) -> Dict:
        """
        Extract text from CSV file

        Args:
            file_path: Path to CSV file

        Returns:
            Dict with formatted text and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            rows = None
            headers = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, newline='') as f:
                        reader = csv_module.reader(f)
                        rows = list(reader)
                    break
                except (UnicodeDecodeError, csv_module.Error):
                    continue

            if rows is None or len(rows) == 0:
                raise Exception("Could not parse CSV file or file is empty")

            # First row is typically headers
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []

            # Format as text
            text_parts = []

            # Add headers
            if headers:
                text_parts.append("Headers: " + " | ".join(headers))
                text_parts.append("-" * 80)

            # Add rows
            for row in data_rows:
                text_parts.append(" | ".join(str(cell) for cell in row))

            combined_text = '\n'.join(text_parts)

            metadata = {
                'row_count': len(rows),
                'column_count': len(headers),
                'headers': headers,
                'has_headers': len(headers) > 0,
                'encoding': encoding
            }

            return {
                'text': combined_text,
                'metadata': metadata,
                'structure': {
                    'headers': headers,
                    'data': data_rows
                },
                'extraction_method': 'csv'
            }

        except Exception as e:
            raise Exception(f"Failed to extract CSV: {str(e)}")

    def extract_metadata_from_file(self, file_path: str) -> Dict:
        """
        Extract filesystem metadata

        Args:
            file_path: Path to file

        Returns:
            Dict with file metadata
        """
        try:
            stat = os.stat(file_path)

            return {
                'file_size': stat.st_size,
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat()
            }
        except Exception as e:
            print(f"Could not extract file metadata: {e}")
            return {}
